"""Shared fixtures for doc-aggregator tests."""

from __future__ import annotations

from pathlib import Path

import cv2
import fitz
import numpy as np
import pytest
from docx import Document


@pytest.fixture()
def tmp_input_dir(tmp_path: Path) -> Path:
    path = tmp_path / "input"
    path.mkdir(parents=True, exist_ok=True)
    return path


def create_txt(path: Path, content: str, encoding: str = "utf-8") -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding=encoding)
    return path


def create_docx(path: Path, heading: str, body: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    doc.save(path)
    return path


def create_text_pdf(path: Path, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), text)
    pdf.save(path)
    pdf.close()
    return path


def create_blank_pdf(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(path)
    pdf.close()
    return path


def create_image(path: Path, text: str = "Fixture OCR") -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = np.zeros((120, 360, 3), dtype=np.uint8)
    cv2.putText(
        image,
        text,
        (10, 70),
        cv2.FONT_HERSHEY_SIMPLEX,
        1.0,
        (255, 255, 255),
        2,
    )
    cv2.imwrite(str(path), image)
    return path


@pytest.fixture()
def generated_small_corpus(tmp_path: Path) -> Path:
    """Runtime-generated mini corpus for parameterized integration tests."""
    root = tmp_path / "fixtures"
    create_txt(root / "small" / "txt" / "runtime_utf8.txt", "runtime text fixture")
    create_docx(root / "small" / "docx" / "runtime.docx", "Runtime", "docx fixture")
    create_text_pdf(root / "small" / "pdf" / "runtime_text.pdf", "runtime pdf text")
    create_blank_pdf(root / "small" / "pdf" / "runtime_scan.pdf")
    create_image(root / "small" / "images" / "runtime.png")
    return root


@pytest.fixture(params=["txt", "docx", "pdf", "image"])
def runtime_input_file(request: pytest.FixtureRequest, tmp_path: Path) -> Path:
    """Parametric input fixture across supported formats."""
    kind = request.param
    if kind == "txt":
        return create_txt(tmp_path / "input.txt", "txt fixture")
    if kind == "docx":
        return create_docx(tmp_path / "input.docx", "Doc", "docx fixture body")
    if kind == "pdf":
        return create_text_pdf(tmp_path / "input.pdf", "pdf fixture")
    return create_image(tmp_path / "input.png")

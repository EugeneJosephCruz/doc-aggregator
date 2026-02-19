"""Generate deterministic binary fixture corpus for tests."""

from __future__ import annotations

from pathlib import Path

import cv2
import fitz
import numpy as np
from docx import Document


def main() -> None:
    root = Path(__file__).resolve().parent
    small = root / "small"
    edge = root / "edge"

    # Small TXT
    (small / "txt").mkdir(parents=True, exist_ok=True)
    (small / "txt" / "utf8_sample.txt").write_text(
        "This is a UTF-8 fixture.\nSecond line.",
        encoding="utf-8",
    )
    (small / "txt" / "ascii_sample.txt").write_text(
        "Simple ASCII fixture for fast tests.",
        encoding="ascii",
    )

    # Small DOCX
    (small / "docx").mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Fixture DOCX", level=1)
    doc.add_paragraph("This is a sample docx fixture paragraph.")
    doc.save(small / "docx" / "sample.docx")

    # Landscape DOCX
    landscape = Document()
    landscape.add_heading("Landscape Fixture", level=1)
    landscape.add_paragraph("Landscape placeholder content.")
    landscape.save(small / "docx" / "landscape_sample.docx")

    # Small PDFs
    (small / "pdf").mkdir(parents=True, exist_ok=True)
    text_pdf = fitz.open()
    page = text_pdf.new_page()
    page.insert_text((72, 72), "Hello from text-native PDF fixture")
    text_pdf.save(small / "pdf" / "text_native.pdf")
    text_pdf.close()

    scanned_pdf = fitz.open()
    scanned_pdf.new_page()  # blank page to trigger OCR path
    scanned_pdf.save(small / "pdf" / "scan_like.pdf")
    scanned_pdf.close()

    # Small images
    (small / "images").mkdir(parents=True, exist_ok=True)
    image = np.zeros((120, 360, 3), dtype=np.uint8)
    cv2.putText(image, "Fixture OCR", (10, 70), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (255, 255, 255), 2)
    cv2.imwrite(str(small / "images" / "clean.png"), image)
    cv2.imwrite(str(small / "images" / "clean.jpg"), image)
    cv2.imwrite(str(small / "images" / "clean.tiff"), image)

    # Corrupted samples
    (edge / "corrupted").mkdir(parents=True, exist_ok=True)
    (edge / "corrupted" / "broken.pdf").write_bytes(b"not a real pdf")
    (edge / "corrupted" / "broken.docx").write_bytes(b"not a real docx zip")
    (edge / "corrupted" / "broken.png").write_bytes(b"not a real image")

    # Oversized samples
    (edge / "oversized").mkdir(parents=True, exist_ok=True)
    (edge / "oversized" / "large.txt").write_text("x" * 2_000_000, encoding="utf-8")

    # Duplicate basenames
    dup_a = edge / "duplicate_names" / "a"
    dup_b = edge / "duplicate_names" / "b"
    dup_a.mkdir(parents=True, exist_ok=True)
    dup_b.mkdir(parents=True, exist_ok=True)
    (dup_a / "same.txt").write_text("a", encoding="utf-8")
    (dup_b / "same.txt").write_text("b", encoding="utf-8")

    # Symlink cases (best effort)
    symlink_root = edge / "symlink_cases"
    symlink_root.mkdir(parents=True, exist_ok=True)
    target = symlink_root / "target.txt"
    target.write_text("target", encoding="utf-8")
    link = symlink_root / "target_link.txt"
    if not link.exists():
        try:
            link.symlink_to(target)
        except OSError:
            pass

    print(f"Generated fixture corpus in: {root}")


if __name__ == "__main__":
    main()

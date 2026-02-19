"""DOCX ingestion preserving source formatting via docxcompose."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docxcompose.composer import Composer


def process_docx_file(
    source_path: Path,
    display_name: str,
    segment_path: Path,
    logger,
) -> dict[str, bool]:
    """Wrap source docx with Heading 1 header and preserve content/styles."""
    segment_path.parent.mkdir(parents=True, exist_ok=True)

    wrapper = Document()
    wrapper.add_heading(display_name, level=1)
    wrapper.add_section(WD_SECTION_START.NEW_PAGE)

    composer = Composer(wrapper)
    composer.append(Document(str(source_path)))
    composer.save(str(segment_path))
    logger.info("Merged docx source %s into segment %s", source_path, segment_path)
    return {"ocr_used": False}

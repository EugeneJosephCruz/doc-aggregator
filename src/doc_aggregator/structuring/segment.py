"""Per-source segment docx builders."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START


def create_text_segment(
    display_name: str,
    text: str,
    segment_path: Path,
    *,
    landscape: bool = False,
) -> Path:
    """Build a standalone segment docx with heading and text body."""
    segment_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(display_name, level=1)

    if text.strip():
        for line in text.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("")

    if landscape:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.save(str(segment_path))
    return segment_path

"""Final segment composer."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docxcompose.composer import Composer

from doc_aggregator.config import AggregatorConfig


class FinalComposer:
    """Merge per-file segment docs into a single master document."""

    def __init__(self, config: AggregatorConfig) -> None:
        self.config = config

    def merge_all(self, segments: list[Path]) -> Document:
        if not segments:
            doc = Document()
            doc.add_heading("Table of Contents", level=1)
            doc.add_paragraph("No supported files were processed.")
            return doc

        master = Document(str(segments[0]))
        composer = Composer(master)
        for segment in segments[1:]:
            composer.append(Document(str(segment)))
        return master

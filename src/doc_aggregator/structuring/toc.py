"""TOC insertion helpers for python-docx."""

from __future__ import annotations

from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def insert_toc(document: Document) -> None:
    """Insert a clickable TOC field at the start of the document."""
    toc_heading = document.add_paragraph("Table of Contents")
    try:
        toc_heading.style = "TOC Heading"
    except Exception:
        # Fallback style only when TOC Heading doesn't exist in template.
        toc_heading.style = "Heading 1"

    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Update fields to generate the table of contents."

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)

    body = document._element.body
    body.insert(0, paragraph._p)
    body.insert(0, toc_heading._p)

    _enable_update_fields_on_open(document)


def _enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    existing = settings.xpath("./w:updateFields")
    if existing:
        existing[0].set(qn("w:val"), "true")
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
